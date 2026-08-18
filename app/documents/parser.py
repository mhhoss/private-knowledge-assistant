"""PDF and DOCX text extraction.

Extracts raw text only. Normalization and chunking happen in `processor.py`; this module
does not know about languages, scripts, or the pipeline stages around it.

PDF extraction shells out to poppler's `pdftotext -layout`, not the `pypdf` library
this project used previously. `pypdf`'s ToUnicode/CMap handling was measured — against
real Persian PDFs produced by LibreOffice and Chrome, not synthetic ones — to recover as
little as ~11% of the true text, both mirrored and outright corrupted (glyph names
leaking into the output as literal text), regardless of its "layout" vs "plain" mode.
`pdftotext -layout` was the strongest of six extractors benchmarked against the same
corpus (pypdf, PyMuPDF, pypdfium2, pdfplumber, pdfminer.six, pdfmux): 81% average
character fidelity and correct RTL/logical reading order on every real document tested,
versus pypdf's 37%, with no per-seat license (see ARCHITECTURE.md's PDF extraction
decision record for the full comparison). The trade-off is a runtime dependency on the
`poppler-utils` system package rather than a pure-Python one — see README.md's Setup
section for the per-platform install step.
"""

from __future__ import annotations

import io
import subprocess
import tempfile
from pathlib import Path

import docx
from docx.table import Table

# Real PDFs in this project's ingestion path are small, single-purpose documents, not
# huge scanned books; anything still running after this is treated as poppler hanging
# rather than legitimately still working, so the file fails instead of blocking a
# request indefinitely.
_PDFTOTEXT_TIMEOUT_SECONDS = 30

_PDFTOTEXT_NOT_FOUND_MESSAGE = (
    "Could not read PDF: the 'pdftotext' command (from poppler-utils) is not "
    "installed or not on PATH. See README.md's Setup section."
)


class ParsingError(RuntimeError):
    """A file could not be read as valid PDF/DOCX, regardless of the reason."""


# Pages/paragraphs join with a blank line so `processor.chunk_text` treats each as a
# paragraph boundary rather than run-on text.
_BLOCK_SEP = "\n\n"


def extract_text(*, file_type: str, content: bytes) -> str:
    """Extract raw text for a file already identified as `file_type` by the loader."""
    if file_type == "pdf":
        return _extract_pdf(content)
    if file_type == "docx":
        return _extract_docx(content)
    raise ValueError(f"Unsupported file_type: {file_type!r}")


def _extract_pdf(content: bytes) -> str:
    # `pdftotext` reads from a real path, not stdin-as-PDF; `delete=False` + a manual
    # `finally` unlink (rather than the usual context-manager auto-delete) because an
    # open, still-locked file handle can't reliably be reopened by a child process on
    # every platform this runs on.
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(content)
        tmp_path = tmp.name

    try:
        try:
            result = subprocess.run(
                ["pdftotext", "-layout", tmp_path, "-"],
                capture_output=True,
                text=True,
                timeout=_PDFTOTEXT_TIMEOUT_SECONDS,
                check=False,
            )
        except FileNotFoundError as error:
            raise ParsingError(_PDFTOTEXT_NOT_FOUND_MESSAGE) from error
        except subprocess.TimeoutExpired as error:
            raise ParsingError(
                f"Could not read PDF: pdftotext did not finish within "
                f"{_PDFTOTEXT_TIMEOUT_SECONDS}s"
            ) from error
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    if result.returncode != 0:
        raise ParsingError(
            f"Could not read PDF: pdftotext exited with status {result.returncode}: "
            f"{result.stderr.strip()}"
        )
    pages = result.stdout.split("\x0c")  # pdftotext separates pages with a form feed
    return _BLOCK_SEP.join(page.strip() for page in pages if page.strip())


def _extract_docx(content: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        tables = [_render_table(table) for table in document.tables]
    except Exception as error:
        raise ParsingError(f"Could not read DOCX: {error}") from error
    return _BLOCK_SEP.join(block for block in (*paragraphs, *tables) if block)


def _render_table(table: Table) -> str:
    """Render a table as one row per line, cells tab-separated."""
    rows = ("\t".join(cell.text for cell in row.cells) for row in table.rows)
    return "\n".join(row for row in rows if row.strip())
