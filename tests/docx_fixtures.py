"""Real DOCX fixtures via python-docx's own writer.

Unlike PDF, no byte-level trick is needed: python-docx builds a genuine .docx, and
`w:t` run text is always stored in logical reading order regardless of script — DOCX has
no equivalent of the PDF visual-order bug.
"""

from __future__ import annotations

import io

import docx


def build_docx(
    paragraphs: list[str], *, tables: list[list[list[str]]] | None = None
) -> bytes:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    for table_rows in tables or []:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row, cells in zip(table.rows, table_rows, strict=True):
            for cell, text in zip(row.cells, cells, strict=True):
                cell.text = text

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
